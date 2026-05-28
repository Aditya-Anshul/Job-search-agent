"""Resume generator — dynamically builds tailored Word resumes (.docx) per job."""

import os
import json
from pathlib import Path
import docx
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from utils.logger import logger, clean_json
from llm import get_llm


async def generate_tailored_resume(
    job_title: str,
    job_description: str,
    profile: dict,
    output_path: str
) -> bool:
    """Dynamically tailors and generates a professional .docx resume matching target job.

    First tries to leverage the existing resume template ('resume/uploads/resume.docx')
    by reading its paragraphs, asking the LLM to selectively edit bullet points,
    skills categories, or summaries, and saving the customized copy to disk.

    If the template is missing, falls back to generating a styled document from scratch.
    """
    logger.info(f"Generating tailored resume for '{job_title}' -> {output_path}")

    # --- Strategy 1: Leverage and Modify Existing .docx Resume Template ---
    existing_docx_path = "resume/uploads/resume.docx"
    if os.path.exists(existing_docx_path):
        try:
            logger.info(f"Leveraging existing resume template from {existing_docx_path}")
            doc = Document(existing_docx_path)
            
            # Map paragraphs by index
            paragraphs_data = {}
            for idx, p in enumerate(doc.paragraphs):
                text = p.text.strip()
                if text:
                    paragraphs_data[str(idx)] = text
            
            # LLM Prompt for editing specific paragraphs
            prompt = (
                f"You are a premium career services resume builder.\n"
                f"Your task is to tailor a candidate's existing resume to match a specific job description by updating individual paragraphs.\n"
                f"You must keep all personal contact details, formatting, structure, education, and dates exactly as they are. "
                f"Only rewrite the achievement bullets under 'PROFESSIONAL EXPERIENCE', 'TECHNICAL SKILLS', 'KEY ACHIEVEMENTS', or projects to align perfectly with the target job's keywords and requirements.\n"
                f"CRITICAL INSTRUCTION: Tailor the existing resume according to the job description WITHOUT changing or updating the core of it. The heart and soul (the actual core truth of the achievements and experience) should remain exactly as it is, but it should be updated, reframed, and injected with keywords to match the job in order to get a high ATS score.\n\n"
                f"Target Job Title: {job_title}\n"
                f"Target Job Description:\n{job_description[:4000]}\n\n"
                f"Existing Resume Paragraphs (indexed):\n{json.dumps(paragraphs_data, indent=2)}\n\n"
                f"Select which paragraph indices to rewrite. Return ONLY a valid JSON object mapping paragraph index strings (e.g. \"6\") to their new tailored texts. Do NOT include markdown backticks or any conversation."
            )
            
            llm = get_llm()
            resp = await llm.complete(prompt)
            clean = clean_json(resp.content)
            updates = json.loads(clean)
            
            if isinstance(updates, dict):
                logger.info(f"Tailored resume updates received from LLM for {len(updates)} paragraphs.")
                for idx_str, new_text in updates.items():
                    try:
                        idx = int(idx_str)
                        if 0 <= idx < len(doc.paragraphs):
                            # Replace paragraph text while preserving style
                            doc.paragraphs[idx].text = new_text
                            logger.debug(f"Updated paragraph {idx} -> {new_text[:50]}...")
                    except Exception as e:
                        logger.warning(f"Failed to update paragraph {idx_str}: {e}")
                
                # Save to output path
                output_dir = os.path.dirname(output_path)
                if output_dir:
                    Path(output_dir).mkdir(parents=True, exist_ok=True)
                doc.save(output_path)
                logger.success(f"Tailored resume successfully created by modifying existing template: {output_path}")
                return True
        except Exception as e:
            logger.error(f"Failed to leverage existing resume template: {e}. Falling back to from-scratch generation.")

    # --- Strategy 2: Fallback to Generating Styled Document From Scratch ---
    logger.info("Generating professional tailored resume from scratch.")
    prompt = (
        f"You are a premium career services resume builder.\n"
        f"Your task is to tailor a candidate's resume content to match a specific job description.\n\n"
        f"Candidate Name: {profile.get('name', 'Candidate')}\n"
        f"Current Role: {profile.get('current_role', 'Engineer')}\n"
        f"Candidate Summary: {profile.get('summary', '')}\n"
        f"Candidate Skills: {', '.join(profile.get('skills', []))}\n\n"
        f"Target Job Title: {job_title}\n"
        f"Target Job Description:\n{job_description[:4000]}\n\n"
        f"CRITICAL INSTRUCTION: Tailor the resume summary, skills, and work experience bullet points to perfectly align with the target job WITHOUT changing the core truth of the candidate's experience. The heart and soul should remain exactly as it is, but updated and optimized with keywords to match the job and get a high ATS score.\n"
        f"Focus on highlighting matching skills and achievements. "
        f"Return ONLY a valid JSON object (no markdown backticks, no conversation) with these exact keys:\n"
        f"  - summary: a tailored 3-4 sentence professional summary\n"
        f"  - highlighted_skills: an array of 8-12 skills that perfectly match the job description\n"
        f"  - experience_bullet_points: an array of 4-6 tailored achievement bullet points\n"
    )

    tailored_data = {
        "summary": profile.get("summary", ""),
        "highlighted_skills": profile.get("skills", [])[:10],
        "experience_bullet_points": [
            "Implemented and managed robust infrastructure for scalable cloud environments.",
            "Automated deployment pipelines to improve delivery speed and system reliability.",
            "Monitored, analyzed, and optimized cloud systems to reduce infrastructure costs."
        ]
    }

    try:
        llm = get_llm()
        resp = await llm.complete(prompt)
        clean = clean_json(resp.content)
        data = json.loads(clean)
        if isinstance(data, dict):
            if "summary" in data:
                tailored_data["summary"] = data["summary"]
            if "highlighted_skills" in data:
                tailored_data["highlighted_skills"] = data["highlighted_skills"]
            if "experience_bullet_points" in data:
                tailored_data["experience_bullet_points"] = data["experience_bullet_points"]
    except Exception as e:
        logger.warning(f"Failed to get tailored resume content from LLM: {e}. Using fallback data.")

    # Build the beautifully formatted Word Document (.docx)
    try:
        doc = Document()

        # Margins (1 inch)
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Header - Candidate Name
        p_name = doc.add_paragraph()
        p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_name = p_name.add_run(profile.get("name", "CANDIDATE").upper())
        run_name.bold = True
        run_name.font.size = Pt(20)
        run_name.font.name = "Calibri"

        # Header - Contact Details
        p_contact = doc.add_paragraph()
        p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_parts = []
        if profile.get("email"):
            contact_parts.append(profile["email"])
        if profile.get("phone"):
            contact_parts.append(profile["phone"])
        contact_parts.append(profile.get("current_role", "Engineer"))
        run_contact = p_contact.add_run(" | ".join(contact_parts))
        run_contact.font.size = Pt(11)
        run_contact.font.name = "Calibri"

        # Section 1: Professional Summary
        doc.add_heading("PROFESSIONAL SUMMARY", level=1)
        p_sum = doc.add_paragraph(tailored_data["summary"])
        p_sum.style.font.name = "Calibri"
        p_sum.style.font.size = Pt(11)

        # Section 2: Core Competencies / Skills
        doc.add_heading("CORE COMPETENCIES", level=1)
        p_skills = doc.add_paragraph()
        p_skills.add_run(", ".join(tailored_data["highlighted_skills"]))
        p_skills.style.font.name = "Calibri"
        p_skills.style.font.size = Pt(11)

        # Section 3: Professional Experience
        doc.add_heading("PROFESSIONAL EXPERIENCE", level=1)
        p_exp_header = doc.add_paragraph()
        run_exp_title = p_exp_header.add_run(f"Senior {profile.get('current_role', 'Engineer')}")
        run_exp_title.bold = True
        run_exp_title.font.size = Pt(12)

        for bullet in tailored_data["experience_bullet_points"]:
            doc.add_paragraph(bullet, style='List Bullet')

        # Section 4: Education & Certifications
        if profile.get("education") or profile.get("certifications"):
            doc.add_heading("EDUCATION & CERTIFICATIONS", level=1)
            if profile.get("education"):
                p_edu = doc.add_paragraph()
                run_edu = p_edu.add_run(profile["education"])
                run_edu.font.size = Pt(11)
            if profile.get("certifications"):
                for cert in profile["certifications"]:
                    doc.add_paragraph(cert, style='List Bullet')

        # Style headings (Calibri, Premium Navy Blue, bold)
        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith("Heading"):
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(14)
                    run.bold = True
                    run.font.color.rgb = docx.shared.RGBColor(0, 51, 102)

        # Save to output path
        output_dir = os.path.dirname(output_path)
        if output_dir:
            Path(output_dir).mkdir(parents=True, exist_ok=True)

        doc.save(output_path)
        logger.success(f"Tailored resume saved successfully at {output_path}")
        return True
    except Exception as e:
        logger.error(f"Failed to generate tailored resume .docx: {e}")
        return False
